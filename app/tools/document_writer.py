"""
app/tools/document_writer.py
Tool — Document Writer

Converts Markdown resume/cover letter content into formatted DOCX files.
"""
from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from loguru import logger

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — DOCX generation disabled")


class DocumentWriter:
    """
    Generates professional DOCX documents from Markdown content.

    Applies clean, ATS-friendly formatting:
    - Standard fonts (Calibri 11pt body, Arial headers)
    - Consistent spacing
    - No tables or columns (ATS compatibility)
    - Clean margins
    """

    OUTPUT_DIR = Path(os.getenv("DATA_DIR", "./data")) / "applications"

    def __init__(self):
        self.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    def save_resume(
        self,
        markdown_content: str,
        candidate_name: str,
        company_name: str,
        version_id: str,
    ) -> Optional[str]:
        """Generate and save resume DOCX. Returns file path."""
        filename = f"resume_{self._slugify(candidate_name)}_{self._slugify(company_name)}_{version_id}.docx"
        filepath = self.OUTPUT_DIR / filename
        return self._write_docx(markdown_content, str(filepath), doc_type="resume")

    def save_cover_letter(
        self,
        markdown_content: str,
        candidate_name: str,
        company_name: str,
        version_id: str,
    ) -> Optional[str]:
        """Generate and save cover letter DOCX. Returns file path."""
        filename = f"cover_letter_{self._slugify(candidate_name)}_{self._slugify(company_name)}_{version_id}.docx"
        filepath = self.OUTPUT_DIR / filename
        return self._write_docx(markdown_content, str(filepath), doc_type="cover_letter")

    def _write_docx(
        self,
        markdown: str,
        filepath: str,
        doc_type: str = "resume"
    ) -> Optional[str]:
        if not DOCX_AVAILABLE:
            logger.warning("[DocWriter] python-docx not available — skipping DOCX generation")
            return None

        try:
            doc = Document()
            self._set_margins(doc)
            self._set_default_style(doc)
            self._parse_markdown_to_docx(doc, markdown, doc_type)
            doc.save(filepath)
            logger.info(f"[DocWriter] Saved {doc_type}: {filepath}")
            return filepath
        except Exception as e:
            logger.error(f"[DocWriter] Failed to write DOCX: {e}")
            return None

    def _set_margins(self, doc: "Document"):
        """Set ATS-friendly 1-inch margins."""
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _set_default_style(self, doc: "Document"):
        """Configure default paragraph style."""
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    def _parse_markdown_to_docx(self, doc: "Document", markdown: str, doc_type: str):
        """
        Parse Markdown and add formatted paragraphs to document.
        Handles: # headings, ## subheadings, - bullets, **bold**, normal paragraphs
        """
        lines = markdown.split("\n")

        for line in lines:
            line = line.rstrip()

            if not line:
                # Empty line — add spacing paragraph
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                continue

            if line.startswith("# "):
                # H1 — Name / title
                p = doc.add_paragraph()
                run = p.add_run(line[2:].strip())
                run.bold = True
                run.font.size = Pt(18)
                run.font.name = "Arial"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)

            elif line.startswith("## "):
                # H2 — Section header
                p = doc.add_paragraph()
                run = p.add_run(line[3:].strip().upper())
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Arial"
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                # Add bottom border effect with separator
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                # Add a horizontal rule paragraph
                hr = doc.add_paragraph("─" * 60)
                hr.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                hr.runs[0].font.size = Pt(8)
                hr.paragraph_format.space_after = Pt(2)

            elif line.startswith("### "):
                # H3 — Job title / company
                p = doc.add_paragraph()
                run = p.add_run(line[4:].strip())
                run.bold = True
                run.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)

            elif line.startswith(("- ", "* ", "• ")):
                # Bullet point
                text = line[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                self._add_formatted_run(p, text)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(1)

            else:
                # Normal paragraph
                p = doc.add_paragraph()
                self._add_formatted_run(p, line)
                p.paragraph_format.space_after = Pt(2)

    def _add_formatted_run(self, paragraph, text: str):
        """Add text with inline bold (**text**) support."""
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def _slugify(self, text: str) -> str:
        """Convert text to filename-safe slug."""
        return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")[:30]
